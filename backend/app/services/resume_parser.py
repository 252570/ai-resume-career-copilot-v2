"""Safe, deterministic PDF/DOCX/TXT text extraction and conservative resume signal parsing."""

from __future__ import annotations

import re
import zipfile
from io import BytesIO
from pathlib import PurePath

from docx import Document
from pypdf import PdfReader

from app.core.errors import ResumeUploadError
from app.schemas.resume import ParsedResumeData

PDF_CONTENT_TYPE = "application/pdf"
DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
TEXT_CONTENT_TYPE = "text/plain"
SUPPORTED_SUFFIXES = {".pdf": PDF_CONTENT_TYPE, ".docx": DOCX_CONTENT_TYPE, ".txt": TEXT_CONTENT_TYPE}

EMAIL_PATTERN = re.compile(r"(?<![\w.+-])[\w.+-]+@[A-Za-z0-9-]+(?:\.[A-Za-z0-9-]+)+")
PHONE_PATTERN = re.compile(r"(?<!\w)(?:\+?\d[\d().\-\s]{6,}\d)(?!\w)")
LINKEDIN_PATTERN = re.compile(r"(?:https?://)?(?:www\.)?linkedin\.com/[A-Za-z0-9_./?=&%-]+", re.IGNORECASE)
GITHUB_PATTERN = re.compile(r"(?:https?://)?(?:www\.)?github\.com/[A-Za-z0-9_./?=&%-]+", re.IGNORECASE)
URL_PATTERN = re.compile(r"https?://[^\s<>()]+", re.IGNORECASE)
SECTION_MARKERS = (
    "education", "academic", "qualification", "experience", "employment", "work history", "skills",
    "projects", "project experience", "certifications", "certificates", "summary", "profile", "professional summary", "portfolio", "links", "contact",
)
# Skill vocabulary for resume + job parsing. Each value is a regex matched with
# re.IGNORECASE against the whole document. Notes on the tricky cases:
#   * Word boundaries (\b) keep sub-string false positives out, e.g. "\bjava\b"
#     will NOT match inside "JavaScript", and "\bsql\b" will NOT match inside
#     "PostgreSQL" or "MySQL".
#   * A skill whose name is also a common English word ("Spring", "Go", "Excel",
#     "Spark") must NEVER be detected from the bare word alone. A resume reading
#     "Spring 2023", "go the extra mile", or "helped excel at" contains no such
#     skill, and crediting one fabricates evidence the document does not support.
#     These require a disambiguating qualifier ("Spring Boot", "Go programming",
#     "Microsoft Excel", "Apache Spark") or an unambiguous list context.
#   * Acronyms that collide with ordinary words when lowercased (ML, AI) use an
#     inline (?-i:...) group to force a case-sensitive match, so prose like
#     "the ml team" or an address such as "ai@example.com" does not match.
#   * Bare single letters like "C" and "R" are intentionally omitted: they are
#     far too noisy to detect reliably from free text.
SKILL_PATTERNS = {
    # --- Programming languages ---
    "Python": r"\bpython\b",
    "JavaScript": r"\bjavascript\b",
    "TypeScript": r"\btypescript\b",
    "Java": r"\bjava\b",
    "C++": r"\bc\+\+",
    "C#": r"\bc#|\bc sharp\b",
    "Go": r"\bgolang\b|(?-i:\bGo\b)(?=[\s,/|)\].]*(?:programming|language|lang\b))|(?-i:\bGo\b)(?=\s*[,/|]\s*(?-i:[A-Z]))",
    "Rust": r"\brust\b",
    "Kotlin": r"\bkotlin\b",
    "Swift": r"\bswift\b",
    "PHP": r"\bphp\b",
    "Ruby": r"\bruby\b",
    "Scala": r"\bscala\b",
    "Perl": r"\bperl\b",
    "MATLAB": r"\bmatlab\b",
    # --- Web / frontend ---
    "HTML": r"\bhtml5?\b",
    "CSS": r"\bcss3?\b",
    "React": r"\breact(?:\.js)?\b",
    "React Native": r"\breact native\b",
    "Next.js": r"\bnext\.?js\b",
    "Angular": r"\bangular(?:js)?\b",
    "Vue.js": r"\bvue(?:\.js)?\b",
    "Node.js": r"\bnode\.?js\b|\bnodejs\b",
    "Express": r"\bexpress(?:\.?js)?\b",
    "Tailwind CSS": r"\btailwind(?: ?css)?\b",
    "Bootstrap": r"\bbootstrap\b",
    # --- Backend frameworks ---
    "FastAPI": r"\bfastapi\b",
    "Django": r"\bdjango\b",
    "Flask": r"\bflask\b",
    "Spring": r"\bspring\s?boot\b|\bspring\s?framework\b|\bspring\s?mvc\b|\bspring\s?security\b|\bspring\s?data\b",
    ".NET": r"\basp\.net\b|\bdotnet\b|(?<![\w.])\.net\b",
    "Flutter": r"\bflutter\b",
    "GraphQL": r"\bgraphql\b",
    "REST APIs": r"\brestful\b|\brest apis?\b",
    # --- Data / ML ---
    "Machine Learning": r"\bmachine\s?learning\b|\bartificial\s?intelligence\b|(?-i:\bML\b)|(?-i:\bAI\b)",
    "Deep Learning": r"\bdeep learning\b",
    "NLP": r"\bnlp\b|\bnatural language processing\b",
    "Computer Vision": r"\bcomputer vision\b",
    "TensorFlow": r"\btensorflow\b",
    "PyTorch": r"\bpytorch\b",
    "scikit-learn": r"\bscikit[- ]?learn\b|\bsklearn\b",
    "Keras": r"\bkeras\b",
    "Pandas": r"\bpandas\b",
    "NumPy": r"\bnumpy\b",
    "Matplotlib": r"\bmatplotlib\b",
    "OpenCV": r"\bopencv\b",
    "Apache Spark": r"\bapache\s?spark\b|\bpyspark\b|\bspark\b(?=\s*(?:sql|streaming|mllib))",
    "Hadoop": r"\bhadoop\b",
    "Tableau": r"\btableau\b",
    "Power BI": r"\bpower bi\b",
    "Excel": r"\bmicrosoft\s?excel\b|\bms\s?excel\b|\bexcel\b(?=\s*(?:spreadsheet|macros|vba|pivot))|(?<=\bin\s)\bexcel\b",
    # --- Databases ---
    "SQL": r"\bsql\b",
    "PostgreSQL": r"\bpostgres(?:ql)?\b",
    "MySQL": r"\bmysql\b",
    "MongoDB": r"\bmongodb\b|\bmongo\b",
    "Redis": r"\bredis\b",
    "SQLite": r"\bsqlite\b",
    "Oracle": r"\boracle\b",
    "Cassandra": r"\bcassandra\b",
    "DynamoDB": r"\bdynamodb\b",
    "Elasticsearch": r"\belasticsearch\b",
    "Firebase": r"\bfirebase\b",
    # --- Cloud / DevOps ---
    "AWS": r"\baws\b|\bamazon web services\b",
    "Azure": r"\bazure\b",
    "GCP": r"\bgcp\b|\bgoogle cloud\b",
    "Docker": r"\bdocker\b",
    "Kubernetes": r"\bkubernetes\b|\bk8s\b",
    "Terraform": r"\bterraform\b",
    "Jenkins": r"\bjenkins\b",
    "Ansible": r"\bansible\b",
    "CI/CD": r"\bci/?cd\b",
    "Linux": r"\blinux\b",
    "Nginx": r"\bnginx\b",
    "Kafka": r"\bkafka\b",
    "Microservices": r"\bmicroservices?\b",
    # --- Tools / practices ---
    "Git": r"\bgit\b",
    "Jira": r"\bjira\b",
    "Agile": r"\bagile\b",
    "Scrum": r"\bscrum\b",
    "Figma": r"\bfigma\b",
    "Postman": r"\bpostman\b",
}


def validate_and_detect(filename: str | None, content: bytes) -> tuple[str, str]:
    """Require an allowed suffix and matching lightweight file-signature validation."""
    suffix = PurePath(filename or "").suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        raise ResumeUploadError("Unsupported file type. Upload a PDF, DOCX, or TXT resume.")
    if not content:
        raise ResumeUploadError("The uploaded resume is empty.")
    if suffix == ".pdf" and not content.startswith(b"%PDF-"):
        raise ResumeUploadError("The uploaded file does not contain a valid PDF signature.")
    if suffix == ".docx":
        if not zipfile.is_zipfile(BytesIO(content)):
            raise ResumeUploadError("The uploaded file does not contain a valid DOCX package.")
        try:
            with zipfile.ZipFile(BytesIO(content)) as document_zip:
                if "[Content_Types].xml" not in document_zip.namelist() or "word/document.xml" not in document_zip.namelist():
                    raise ResumeUploadError("The uploaded file does not contain a valid DOCX document.")
        except zipfile.BadZipFile as exc:
            raise ResumeUploadError("The uploaded file does not contain a valid DOCX package.") from exc
    if suffix == ".txt":
        _decode_text(content)
    return suffix, SUPPORTED_SUFFIXES[suffix]


def extract_resume_text(content: bytes, suffix: str) -> str:
    """Extract normalized readable text or raise a controlled validation error."""
    try:
        if suffix == ".pdf":
            reader = PdfReader(BytesIO(content), strict=False)
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
        elif suffix == ".docx":
            document = Document(BytesIO(content))
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
            text = "\n".join(part for part in (text, table_text) if part)
        elif suffix == ".txt":
            text = _decode_text(content)
        else:
            raise ResumeUploadError("Unsupported file type. Upload a PDF, DOCX, or TXT resume.")
    except ResumeUploadError:
        raise
    except Exception as exc:
        raise ResumeUploadError("The uploaded document could not be read. Please upload a readable resume.") from exc

    normalized = _normalize_text(text)
    if not normalized:
        raise ResumeUploadError("The uploaded document does not contain readable text.")
    return normalized


def parse_resume_text(text: str) -> ParsedResumeData:
    """Extract only observable signals using regular expressions and section boundaries."""
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    email = _first_match(EMAIL_PATTERN, text)
    phone = _first_match(PHONE_PATTERN, text)
    return ParsedResumeData(
        candidate_name=_candidate_name(lines),
        email=email,
        phone=phone,
        linkedin=_normalise_url(_first_match(LINKEDIN_PATTERN, text)),
        github=_normalise_url(_first_match(GITHUB_PATTERN, text)),
        summary=_section_lines(lines, ("summary", "profile", "professional summary")),
        skills=[name for name, pattern in SKILL_PATTERNS.items() if re.search(pattern, text, re.IGNORECASE)],
        education=_section_lines(lines, ("education", "academic", "qualification")),
        experience=_section_lines(lines, ("experience", "employment", "work history")),
        projects=_section_lines(lines, ("projects", "project experience")),
        certifications=_section_lines(lines, ("certifications", "certificates")),
        links=_unique_links(text),
    )


def _decode_text(content: bytes) -> str:
    if content.startswith((b"\xff\xfe", b"\xfe\xff")):
        try:
            return content.decode("utf-16")
        except UnicodeDecodeError as exc:
            raise ResumeUploadError("The TXT file uses an unreadable text encoding.") from exc
    try:
        return content.decode("utf-8-sig")
    except UnicodeDecodeError as exc:
        raise ResumeUploadError("The TXT file must use UTF-8 or UTF-16 text encoding.") from exc


def _normalize_text(text: str) -> str:
    lines = [re.sub(r"[\t ]+", " ", line).strip() for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    compact_lines: list[str] = []
    previous_blank = False
    for line in lines:
        if line:
            compact_lines.append(line)
            previous_blank = False
        elif not previous_blank:
            compact_lines.append("")
            previous_blank = True
    return "\n".join(compact_lines).strip()


def _first_match(pattern: re.Pattern[str], text: str) -> str | None:
    match = pattern.search(text)
    return match.group(0).rstrip(".,;:)]") if match else None


def _normalise_url(value: str | None) -> str | None:
    if not value:
        return None
    return value if value.startswith(("http://", "https://")) else f"https://{value}"


def _unique_links(text: str) -> list[str]:
    links: list[str] = []
    for match in URL_PATTERN.finditer(text):
        value = match.group(0).rstrip(".,;:)")
        if value not in links:
            links.append(value)
        if len(links) == 12:
            break
    return links


def _candidate_name(lines: list[str]) -> str | None:
    for line in lines[:8]:
        lower = line.lower()
        words = line.split()
        if (
            2 <= len(words) <= 5
            and 3 <= len(line) <= 80
            and not re.search(r"\d|@|https?://|linkedin|github", lower)
            and not any(marker in lower for marker in SECTION_MARKERS)
        ):
            return line
    return None


def _section_lines(lines: list[str], headings: tuple[str, ...]) -> list[str]:
    start: int | None = None
    for index, line in enumerate(lines):
        normalized = re.sub(r"[:\-–—]+$", "", line.lower()).strip()
        if any(normalized == heading or normalized.startswith(f"{heading} ") for heading in headings):
            start = index + 1
            break
    if start is None:
        return []
    extracted: list[str] = []
    for line in lines[start:]:
        normalized = re.sub(r"[:\-–—]+$", "", line.lower()).strip()
        if any(normalized == marker or normalized.startswith(f"{marker} ") for marker in SECTION_MARKERS):
            break
        extracted.append(line)
        if len(extracted) == 8:
            break
    return extracted
