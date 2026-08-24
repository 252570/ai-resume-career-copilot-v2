"""End-to-end API tests for safe Phase 3 resume upload, parsing, storage, and retrieval."""

from __future__ import annotations

from io import BytesIO

import pytest
from docx import Document
from fastapi.testclient import TestClient
from sqlalchemy.orm import Session

from app.api.v1.routes.resumes import get_resume_storage
from app.db.session import get_db_session
from app.main import app
from app.services.resume_parser import DOCX_CONTENT_TYPE, PDF_CONTENT_TYPE, TEXT_CONTENT_TYPE, parse_resume_text
from app.services.resume_storage import ResumeStorage


@pytest.fixture
def client(db_session: Session, tmp_path):
    def override_session():
        yield db_session

    app.dependency_overrides[get_db_session] = override_session
    app.dependency_overrides[get_resume_storage] = lambda: ResumeStorage(tmp_path / "resumes")
    with TestClient(app) as test_client:
        yield test_client
    app.dependency_overrides.clear()


def _pdf_with_text(text: str) -> bytes:
    escaped = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    stream = f"BT /F1 16 Tf 72 720 Td ({escaped}) Tj ET".encode()
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>",
        b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    output = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, body in enumerate(objects, start=1):
        offsets.append(len(output))
        output.extend(f"{index} 0 obj\n".encode() + body + b"\nendobj\n")
    xref_start = len(output)
    output.extend(f"xref\n0 {len(objects) + 1}\n0000000000 65535 f \n".encode())
    output.extend(b"".join(f"{offset:010d} 00000 n \n".encode() for offset in offsets[1:]))
    output.extend(f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_start}\n%%EOF".encode())
    return bytes(output)


def _docx_with_text() -> bytes:
    document = Document()
    document.add_paragraph("Avery Example")
    document.add_paragraph("avery@example.com")
    document.add_paragraph("Skills")
    document.add_paragraph("Python, FastAPI, PostgreSQL")
    output = BytesIO()
    document.save(output)
    return output.getvalue()


@pytest.mark.parametrize(
    ("filename", "content", "content_type"),
    [
        ("candidate.pdf", _pdf_with_text("Avery Example avery@example.com Python"), PDF_CONTENT_TYPE),
        ("candidate.docx", _docx_with_text(), DOCX_CONTENT_TYPE),
        ("candidate.txt", b"Avery Example\navery@example.com\nSkills\nPython, FastAPI", TEXT_CONTENT_TYPE),
    ],
    # Explicit ids are required, not cosmetic. Without them pytest derives the id from the
    # raw fixture bytes, and a ~36 KB DOCX expands to a ~105 KB id. pytest exports that id
    # in the PYTEST_CURRENT_TEST environment variable, which exceeds the 32,767-character
    # Windows limit and raises ValueError during setup, so the case never runs at all.
    ids=["pdf", "docx", "txt"],
)
def test_supported_resumes_upload_and_retrieve(client: TestClient, filename: str, content: bytes, content_type: str) -> None:
    upload = client.post("/api/v1/resumes/upload", files={"file": (filename, content, content_type)})

    assert upload.status_code == 201, upload.text
    payload = upload.json()
    assert payload["filename"] == filename
    assert payload["content_type"] == content_type
    assert payload["status"] == "parsed"
    assert "id" in payload

    retrieved = client.get(f"/api/v1/resumes/{payload['id']}")
    assert retrieved.status_code == 200
    assert retrieved.json()["id"] == payload["id"]
    assert "storage_key" not in retrieved.json()


def test_rejects_unsupported_large_and_unreadable_files(client: TestClient) -> None:
    unsupported = client.post("/api/v1/resumes/upload", files={"file": ("resume.png", b"not an image", "image/png")})
    too_large = client.post("/api/v1/resumes/upload", files={"file": ("resume.txt", b"a" * (5 * 1024 * 1024 + 1), TEXT_CONTENT_TYPE)})
    empty = client.post("/api/v1/resumes/upload", files={"file": ("resume.txt", b"", TEXT_CONTENT_TYPE)})
    unreadable = client.post("/api/v1/resumes/upload", files={"file": ("resume.pdf", b"%PDF-not-readable", PDF_CONTENT_TYPE)})

    assert unsupported.status_code == 400
    assert too_large.status_code == 413
    assert empty.status_code == 400
    assert unreadable.status_code == 400


def test_deterministic_contact_url_and_skill_extraction() -> None:
    parsed = parse_resume_text(
        "Avery Example\navery@example.com\n+1 (555) 222-3333\n"
        "linkedin.com/in/avery\ngithub.com/avery\nSkills\nPython, FastAPI, PostgreSQL\n"
        "Education\nExample University\nExperience\nSoftware Engineer"
    )

    assert parsed.email == "avery@example.com"
    assert parsed.linkedin == "https://linkedin.com/in/avery"
    assert parsed.github == "https://github.com/avery"
    assert {"Python", "FastAPI", "PostgreSQL"}.issubset(parsed.skills)
    assert parsed.education == ["Example University"]
    assert parsed.experience == ["Software Engineer"]


def test_deterministic_resume_sections_and_links() -> None:
    parsed = parse_resume_text(
        "Avery Example\nSummary\nBackend engineer focused on reliable APIs.\n"
        "Projects\nCareer Copilot API\nCertifications\nAWS Certified Cloud Practitioner\n"
        "Portfolio https://example.com/avery\n"
    )

    assert parsed.summary == ["Backend engineer focused on reliable APIs."]
    assert parsed.projects == ["Career Copilot API"]
    assert parsed.certifications == ["AWS Certified Cloud Practitioner"]
    assert parsed.links == ["https://example.com/avery"]


def test_lists_previously_uploaded_resumes(client: TestClient) -> None:
    client.post(
        "/api/v1/resumes/upload",
        files={"file": ("first.txt", b"Avery Example\nSkills\nPython", TEXT_CONTENT_TYPE)},
    )
    client.post(
        "/api/v1/resumes/upload",
        files={"file": ("second.txt", b"Avery Example\nSkills\nFastAPI", TEXT_CONTENT_TYPE)},
    )

    response = client.get("/api/v1/resumes")
    assert response.status_code == 200
    assert {item["filename"] for item in response.json()} == {"second.txt", "first.txt"}


# --- Skill-detection precision -------------------------------------------------------
# The product boundary in README.md is that parsing never fabricates evidence the source
# document does not contain. Several skill names are also ordinary English words, so a
# bare word match credits a candidate with skills they never claimed, and that error then
# propagates into the match score, ATS gaps, roadmap steps, and project prompts.


@pytest.mark.parametrize(
    ("skill", "text"),
    [
        ("Spring", "B.Tech Computer Science, Spring 2023 - Fall 2026"),
        ("Spring", "Completed a Spring semester internship"),
        ("Go", "Go the extra mile for every customer"),
        ("Go", "Going above and beyond as the go-to person on the team"),
        ("Excel", "Helped the team excel at delivering features on time"),
        ("Excel", "Excellent written and verbal communication skills"),
        ("Apache Spark", "Sparked a redesign of the onboarding flow"),
        ("Machine Learning", "Contact ai@example.com or reach the ml team"),
    ],
    ids=[
        "spring-semester-date",
        "spring-semester-word",
        "go-idiom",
        "go-gerund",
        "excel-verb",
        "excel-excellent",
        "spark-verb",
        "ml-ai-lowercase-prose",
    ],
)
def test_common_words_are_not_reported_as_skills(skill: str, text: str) -> None:
    """A skill must never be inferred from an ordinary English word used as prose."""
    parsed = parse_resume_text(f"Avery Example\nExperience\n{text}\n")

    assert skill not in parsed.skills, f"fabricated {skill!r} from: {text!r}"


@pytest.mark.parametrize(
    ("skill", "text"),
    [
        ("Spring", "Built REST services with Spring Boot"),
        ("Spring", "Experience with Spring MVC and Hibernate"),
        ("Go", "Go programming language"),
        ("Go", "Backend services in Go, Rust, and Python"),
        ("Go", "Proficient in Golang"),
        ("Excel", "Advanced Microsoft Excel and pivot tables"),
        ("Excel", "Excel macros for automated reporting"),
        ("Apache Spark", "Processed batch data with Apache Spark"),
        ("Apache Spark", "Used Spark SQL for ETL pipelines"),
        ("Machine Learning", "Machine Learning coursework and projects"),
        ("Machine Learning", "Applied ML and AI techniques to forecasting"),
        ("Java", "Backend development in Java and Spring Boot"),
    ],
    ids=[
        "spring-boot",
        "spring-mvc",
        "go-language",
        "go-in-list",
        "golang",
        "microsoft-excel",
        "excel-macros",
        "apache-spark",
        "spark-sql",
        "machine-learning",
        "ml-ai-uppercase",
        "java-not-javascript",
    ],
)
def test_genuine_skills_are_still_detected(skill: str, text: str) -> None:
    """Tightening the patterns must not cost real detections."""
    parsed = parse_resume_text(f"Avery Example\nSkills\n{text}\n")

    assert skill in parsed.skills, f"missed {skill!r} in: {text!r}"


def test_substring_skills_do_not_collide() -> None:
    """Related technology names must not match inside one another."""
    parsed = parse_resume_text("Avery Example\nSkills\nJavaScript, PostgreSQL, MySQL\n")

    assert {"JavaScript", "PostgreSQL", "MySQL"}.issubset(parsed.skills)
    assert "Java" not in parsed.skills
    assert "SQL" not in parsed.skills
